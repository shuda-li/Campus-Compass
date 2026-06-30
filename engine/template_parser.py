"""
Word 模板结构解析器。
读取用户上传的 .docx 文件，提取其结构骨架（节/表/占位符），
输出 template_structure JSON，供 LLM prompt 注入和渲染器适配使用。
"""

from docx import Document
import os


def parse_template(docx_path: str) -> dict:
    """
    解析一个 .docx 模板文件，返回结构骨架。
    两遍扫描：先收集原始元素，再合并相邻的 heading+table 对。
    """
    doc = Document(docx_path)
    raw_elements = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is None:
                continue
            text = para.text.strip()
            is_heading = _is_heading(para) or _looks_like_section_title(text)
            raw_elements.append({
                "kind": "paragraph",
                "text": text,
                "is_heading": is_heading,
                "has_placeholder": _has_placeholders(para),
            })
        elif tag == "tbl":
            table = _find_table(doc, element)
            raw_elements.append({
                "kind": "table",
                "table": table,
            })

    # 两遍合并：相邻的 heading + table → 表标题为 heading 文本
    merged = []
    skip_next = False
    for i, elem in enumerate(raw_elements):
        if skip_next:
            skip_next = False
            continue
        if (elem["kind"] == "paragraph" and elem["is_heading"] and elem["text"] and
                i + 1 < len(raw_elements) and raw_elements[i + 1]["kind"] == "table"):
            # heading 后紧跟 table → 合并
            table = raw_elements[i + 1]["table"]
            table_info = _parse_table(table)
            merged.append({
                "kind": "table",
                "title": elem["text"],
                "columns": table_info["columns"],
                "rows": table_info["rows"],
                "has_placeholder": elem["has_placeholder"] or table_info.get("has_placeholder", False),
            })
            skip_next = True
        else:
            merged.append(elem)

    # 生成最终 sections
    sections = []
    order = 0
    current_paragraph = None

    def flush_paragraph():
        nonlocal current_paragraph
        if current_paragraph:
            sections.append(current_paragraph)
            current_paragraph = None

    for elem in merged:
        if elem["kind"] == "table":
            flush_paragraph()
            order += 1
            sections.append({
                "order": order,
                "title": elem.get("title", "表格"),
                "type": "table",
                "columns": elem.get("columns", []),
                "rows": elem.get("rows", []),
                "has_placeholder": elem.get("has_placeholder", False),
            })
        elif elem["kind"] == "paragraph":
            text = elem["text"]
            if not text:
                flush_paragraph()
                continue
            if elem["is_heading"]:
                flush_paragraph()
                order += 1
                current_paragraph = {
                    "order": order,
                    "title": text,
                    "type": "paragraph",
                    "hint": "",
                    "has_placeholder": elem.get("has_placeholder", False),
                }
            elif current_paragraph:
                if not current_paragraph.get("hint"):
                    current_paragraph["hint"] = text[:80]
                if elem.get("has_placeholder"):
                    current_paragraph["has_placeholder"] = True

    flush_paragraph()

    return {
        "source": os.path.basename(docx_path),
        "sections": sections,
    }


def _find_paragraph(doc: Document, oxml_element):
    for para in doc.paragraphs:
        if para._element is oxml_element:
            return para
    return None


def _find_table(doc: Document, oxml_element):
    for table in doc.tables:
        if table._element is oxml_element:
            return table
    return None


def _is_heading(para) -> bool:
    style_name = (para.style.name if para.style else "").lower()
    return any(h in style_name for h in ["heading", "标题", "title"])


def _looks_like_section_title(text: str) -> bool:
    """启发式判断：仅通过中文序号前缀识别标题。"""
    text = text.strip()
    if len(text) > 40:
        return False
    if text.endswith("：") or text.endswith(":"):
        return False
    for prefix in ["一、", "二、", "三、", "四、", "五、", "六、",
                   "1.", "2.", "3.", "4.", "5.", "6.",
                   "第", "（一）", "（二）", "（三）"]:
        if text.startswith(prefix):
            return True
    return False


def _has_placeholders(para) -> bool:
    import re
    return bool(re.search(r"\{[^}]+\}", para.text))


def _parse_table(table) -> dict:
    rows = table.rows
    if len(rows) < 2:
        return {"columns": ["内容"], "rows": []}

    header_cells = [cell.text.strip() for cell in rows[0].cells]
    columns = [h if h else f"列{i+1}" for i, h in enumerate(header_cells)]

    parsed_rows = []
    has_placeholder = False
    for row in rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if not any(cells):
            continue
        label = cells[0] if len(cells) > 0 else ""
        value = cells[1] if len(cells) > 1 else ""

        import re
        if re.search(r"\{[^}]+\}", label) or re.search(r"\{[^}]+\}", value):
            has_placeholder = True

        field = _map_to_plan_field(label)
        parsed_rows.append({"label": label, "value": value, "field": field})

    caption = parsed_rows[0]["label"] if parsed_rows else "表格"
    return {
        "caption": caption[:30],
        "columns": columns,
        "rows": parsed_rows,
        "has_placeholder": has_placeholder,
    }


def _map_to_plan_field(label: str) -> str or None:
    mapping = {
        "活动名称": "activity_topic", "活动主题": "activity_topic",
        "活动标题": "activity_topic", "标题": "activity_topic",
        "主题": "activity_topic",
        "活动目的": "activity_purpose", "活动背景": "activity_purpose",
        "活动意义": "activity_purpose", "目的": "activity_purpose",
        "背景": "activity_purpose",
        "活动时间": "activity_time", "时间": "activity_time",
        "日期": "activity_time",
        "主办单位": "organizer", "主办方": "organizer", "主办": "organizer",
        "承办单位": "host", "承办方": "host", "承办": "host",
        "协办单位": "host",
        "参与人数": "_participants_override", "人数": "_participants_override",
        "预算": "_budget", "经费": "_budget", "总预算": "_budget",
        "地点": "_venue", "场地": "_venue", "活动地点": "_venue",
    }
    label_clean = label.strip().rstrip("：:").rstrip()
    return mapping.get(label_clean, None)


# ═══════════════════════════════════════════════════════════════
#  Markdown 模板解析
# ═══════════════════════════════════════════════════════════════


def parse_markdown_template(md_path: str) -> dict:
    """
    解析 Markdown 模板文件，提取与 parse_template 相同的结构骨架。

    规则：
    - `# ` / `## ` 开头的行 → 节的标题
    - `| xxx | yyy |` 连续多行 → 表格（第一行为表头）
    - 其他非空行 → 属于前一个 paragraph 节的正文
    """
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    raw_elements = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 标题行
        if line.lstrip().startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            raw_elements.append({
                "kind": "paragraph",
                "text": text,
                "is_heading": True,
                "has_placeholder": _md_has_placeholder(text),
                "heading_level": level,
            })
            i += 1
            continue

        # 表格检测：以 | 开头
        if line.strip().startswith("|"):
            table_lines = [line]
            i += 1
            # 收集连续的表格行
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            table_info = _parse_md_table(table_lines)
            raw_elements.append({
                "kind": "table",
                "title": table_info.get("caption", "表格"),
                "columns": table_info["columns"],
                "rows": table_info["rows"],
                "has_placeholder": table_info.get("has_placeholder", False),
            })
            continue

        # 普通段落
        text = line.strip()
        raw_elements.append({
            "kind": "paragraph",
            "text": text,
            "is_heading": False,
            "has_placeholder": _md_has_placeholder(text),
        })
        i += 1

    # 合并相邻的 heading + table（同 Word 逻辑）
    merged = []
    skip = False
    for j, elem in enumerate(raw_elements):
        if skip:
            skip = False
            continue
        if (elem["kind"] == "paragraph" and elem["is_heading"] and
                j + 1 < len(raw_elements) and raw_elements[j + 1]["kind"] == "table"):
            t = raw_elements[j + 1]
            merged.append({
                "kind": "table",
                "title": elem["text"],
                "columns": t["columns"],
                "rows": t["rows"],
                "has_placeholder": elem["has_placeholder"] or t["has_placeholder"],
            })
            skip = True
        else:
            merged.append(elem)

    # 生成 sections
    sections = []
    order = 0
    current = None

    def flush():
        nonlocal current
        if current:
            sections.append(current)
            current = None

    for elem in merged:
        if elem["kind"] == "table":
            flush()
            order += 1
            sections.append({
                "order": order,
                "title": elem.get("title", "表格"),
                "type": "table",
                "columns": elem["columns"],
                "rows": elem["rows"],
                "has_placeholder": elem.get("has_placeholder", False),
            })
        elif elem["kind"] == "paragraph":
            if elem["is_heading"]:
                flush()
                order += 1
                current = {
                    "order": order,
                    "title": elem["text"],
                    "type": "paragraph",
                    "hint": "",
                    "has_placeholder": elem.get("has_placeholder", False),
                }
            elif current:
                if not current.get("hint"):
                    current["hint"] = elem["text"][:80]
                if elem.get("has_placeholder"):
                    current["has_placeholder"] = True

    flush()
    return {"source": os.path.basename(md_path), "sections": sections}


def _md_has_placeholder(text: str) -> bool:
    import re
    return bool(re.search(r"\{[^}]+\}", text))


def _parse_md_table(lines: list) -> dict:
    """解析 Markdown 管道表格。第一行是表头，第二行是分隔符，余下是数据。"""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return {"columns": ["内容"], "rows": []}

    # 跳过第二行（分隔符 |---|---|）
    header = rows[0]
    data_start = 2 if len(rows) > 2 and all(c.replace("-", "").replace(":", "").strip() == "" for c in rows[1]) else 1

    columns = [h if h else f"列{i+1}" for i, h in enumerate(header)]
    parsed_rows = []
    has_placeholder = False

    for row in rows[data_start:]:
        if not any(row):
            continue
        label = row[0] if len(row) > 0 else ""
        value = row[1] if len(row) > 1 else ""
        if _md_has_placeholder(label) or _md_has_placeholder(value):
            has_placeholder = True
        field = _map_to_plan_field(label)
        parsed_rows.append({"label": label, "value": value, "field": field})

    caption = parsed_rows[0]["label"] if parsed_rows else "表格"
    return {
        "caption": caption[:30],
        "columns": columns,
        "rows": parsed_rows,
        "has_placeholder": has_placeholder,
    }


def parse_any_template(path: str) -> dict:
    """统一入口：根据文件扩展名选择解析器。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return parse_template(path)
    elif ext in (".md", ".markdown", ".txt"):
        return parse_markdown_template(path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
