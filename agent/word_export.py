"""
方案导出为 Word 文档 (.docx)

用法:
    from agent.word_export import export_plan_to_docx
    filepath = export_plan_to_docx(plan)
    # → 返回临时文件路径，可直接 send_file
"""

import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── 星穹铁道主题色 ──
NEBULA_PURPLE = RGBColor(0x8B, 0x5C, 0xF6)
AMBER_GOLD = RGBColor(0xF0, 0xC0, 0x60)
DARK_BG = RGBColor(0x0A, 0x0E, 0x1A)
WHITE_TEXT = RGBColor(0xE2, 0xE8, 0xF0)
GRAY_TEXT = RGBColor(0x64, 0x74, 0x8B)


def _set_cell_shading(cell, color):
    """设置单元格背景色。"""
    # python-docx RGBColor 支持 [0][1][2] 索引访问 R/G/B
    r, g, b = color[0], color[1], color[2]
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): f'{r:02X}{g:02X}{b:02X}',
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def _add_styled_paragraph(doc, text: str, style: str = None, bold: bool = False,
                          color: RGBColor = None, size: Pt = None, alignment=None,
                          space_after: Pt = Pt(6)):
    """添加格式化段落。"""
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


def export_plan_to_docx(plan: dict, rooms: list = None) -> str:
    """
    将活动方案导出为 Word 文档。

    Args:
        plan: 活动方案 dict
        rooms: 可选，推荐教室列表

    Returns:
        str: 生成的 .docx 文件路径（临时文件）
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 设置默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ═══════════════════════════════════════
    #  标题
    # ═══════════════════════════════════════
    title = plan.get("activity_topic", "活动策划方案")
    _add_styled_paragraph(doc, title, bold=True, size=Pt(22),
                          color=NEBULA_PURPLE, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(4))

    _add_styled_paragraph(doc, "Campus Compass · 校园活动策划方案", size=Pt(9),
                          color=GRAY_TEXT, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(16))

    # ── 分隔线 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '8B5CF6',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ═══════════════════════════════════════
    #  活动目的
    # ═══════════════════════════════════════
    purpose = plan.get("activity_purpose", "")
    if purpose:
        _add_styled_paragraph(doc, "✦ 活动目的", bold=True, size=Pt(14),
                              color=AMBER_GOLD, space_after=Pt(8))
        _add_styled_paragraph(doc, purpose, size=Pt(11), color=WHITE_TEXT,
                              space_after=Pt(14))

    # ═══════════════════════════════════════
    #  基本信息表
    # ═══════════════════════════════════════
    _add_styled_paragraph(doc, "✦ 基本信息", bold=True, size=Pt(14),
                          color=AMBER_GOLD, space_after=Pt(8))

    info_table = doc.add_table(rows=3, cols=2, style='Table Grid')
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("活动时间", plan.get("activity_time", "待定")),
        ("主办单位", plan.get("organizer", "待定")),
        ("承办单位", plan.get("host", "待定")),
    ]
    for i, (label, value) in enumerate(info_data):
        # Label cell
        cell_l = info_table.cell(i, 0)
        cell_l.text = ""
        p_l = cell_l.paragraphs[0]
        run_l = p_l.add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = AMBER_GOLD
        _set_cell_shading(cell_l, RGBColor(0x11, 0x18, 0x33))
        p_l.paragraph_format.space_before = Pt(4)
        p_l.paragraph_format.space_after = Pt(4)

        # Value cell
        cell_v = info_table.cell(i, 1)
        cell_v.text = ""
        p_v = cell_v.paragraphs[0]
        run_v = p_v.add_run(value)
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = WHITE_TEXT
        _set_cell_shading(cell_v, RGBColor(0x11, 0x18, 0x33))
        p_v.paragraph_format.space_before = Pt(4)
        p_v.paragraph_format.space_after = Pt(4)

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()  # Spacer

    # ═══════════════════════════════════════
    #  活动内容
    # ═══════════════════════════════════════
    content_list = plan.get("activity_content", [])
    if content_list:
        _add_styled_paragraph(doc, f"✦ 活动内容（{len(content_list)} 个环节）",
                              bold=True, size=Pt(14), color=AMBER_GOLD, space_after=Pt(10))

        for i, item in enumerate(content_list):
            phase = item.get("phase", f"环节 {i+1}")
            duration = item.get("duration", "")
            content_text = item.get("content", "")
            host_guide = item.get("host_guide", "")
            interaction = item.get("interaction", "")

            # Phase header
            header = f"环节 {i+1}：{phase}"
            if duration:
                header += f"（{duration}）"
            _add_styled_paragraph(doc, header, bold=True, size=Pt(11),
                                  color=NEBULA_PURPLE, space_after=Pt(4))

            if content_text:
                _add_styled_paragraph(doc, content_text, size=Pt(10),
                                      color=WHITE_TEXT, space_after=Pt(4))

            if host_guide:
                _add_styled_paragraph(doc, f"✦ 引导语：{host_guide}", size=Pt(9),
                                      color=AMBER_GOLD, space_after=Pt(2))

            if interaction:
                _add_styled_paragraph(doc, f"互动方式：{interaction}", size=Pt(9),
                                      color=GRAY_TEXT, space_after=Pt(10))

    # ═══════════════════════════════════════
    #  活动物资
    # ═══════════════════════════════════════
    materials = plan.get("activity_materials", [])
    if materials:
        _add_styled_paragraph(doc, f"✦ 活动物资（{len(materials)} 项）",
                              bold=True, size=Pt(14), color=AMBER_GOLD, space_after=Pt(8))

        mat_table = doc.add_table(rows=len(materials), cols=3, style='Table Grid')
        mat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for j, header in enumerate(["物资名称", "规格", "数量"]):
            cell = mat_table.cell(0, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = AMBER_GOLD
            _set_cell_shading(cell, RGBColor(0x11, 0x18, 0x33))
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

        for i, m in enumerate(materials):
            for j, key in enumerate(["name", "spec", "qty"]):
                cell = mat_table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(m.get(key, "")))
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE_TEXT
                _set_cell_shading(cell, RGBColor(0x11, 0x18, 0x33))
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)

        for row in mat_table.rows:
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(6)
            row.cells[2].width = Cm(3)

    # ═══════════════════════════════════════
    #  推荐教室
    # ═══════════════════════════════════════
    if rooms:
        _add_styled_paragraph(doc, f"✦ 推荐教室（{len(rooms)} 间可选）",
                              bold=True, size=Pt(14), color=AMBER_GOLD, space_after=Pt(8))

        for i, room in enumerate(rooms[:3]):
            rid = room.get("room_id", "?")
            building = room.get("building", "")
            cap = room.get("capacity", "?")
            equip = room.get("equipment", [])
            if isinstance(equip, str):
                try:
                    import json
                    equip = json.loads(equip)
                except Exception:
                    equip = [equip]
            equip_str = " · ".join(equip[:3]) if equip else ""

            rank_mark = "★" if i == 0 else f"  {i+1}"
            line = f"{rank_mark} {rid} · {building}  |  {cap}人"
            if equip_str:
                line += f"  |  {equip_str}"
            _add_styled_paragraph(doc, line, size=Pt(10), color=WHITE_TEXT,
                                  space_after=Pt(4))

    # ── 结尾 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Campus Compass 智能活动策划助手 —")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY_TEXT
    run.italic = True

    # ── 保存到临时文件 ──
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name
